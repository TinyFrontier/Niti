"""Text extraction from CV documents.

Runs synchronously on upload: parsing a resume-sized PDF or DOCX takes tens of
milliseconds, so a queue would cost more than it saves. Only formats we can read
reliably are attempted; anything else is marked `unsupported` and the file is
still stored, so the user never loses an upload to a parser we don't have.

Failures are recorded as short machine-readable codes, never as parser output or
document content — extraction errors end up in the DB and in API responses.
"""

import hashlib
import io
import re
import unicodedata
from dataclasses import dataclass

import pypdf
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.common.enums import CVExtractionStatus

# Formats with a reliable text layer. Legacy .doc is a binary OLE container that
# needs an external converter, so it stays unsupported.
SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx"})

# A real CV always clears this. Falling short means a scan or an image-only PDF.
MIN_TEXT_LENGTH = 100
# Upper bound on stored text, so one pathological document can't bloat the row
# or a later AI prompt. ~50 pages of dense text.
MAX_TEXT_LENGTH = 120_000

# extraction_error values (no document content, safe to show and to log)
ERROR_UNSUPPORTED_FORMAT = "unsupported_format"
ERROR_EMPTY_FILE = "empty_file"
ERROR_ENCRYPTED = "encrypted"
ERROR_PARSE_FAILED = "parse_failed"
ERROR_NO_TEXT_LAYER = "no_text_layer"

_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_SPACES = re.compile(r"[ \t ]+")
_BLANK_LINES = re.compile(r"\n{3,}")


@dataclass
class ExtractionResult:
    status: CVExtractionStatus
    text: str | None = None
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.status is CVExtractionStatus.COMPLETED


def content_hash(data: bytes) -> str:
    """Identifies the exact bytes, so a re-upload of the same file is detectable."""
    return hashlib.sha256(data).hexdigest()


def extract_text(data: bytes, extension: str) -> ExtractionResult:
    """Never raises: any parser problem becomes a failed result with a code."""
    if extension not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(CVExtractionStatus.UNSUPPORTED, error=ERROR_UNSUPPORTED_FORMAT)
    if not data:
        return ExtractionResult(CVExtractionStatus.FAILED, error=ERROR_EMPTY_FILE)

    try:
        raw = _from_pdf(data) if extension == ".pdf" else _from_docx(data)
    except _EncryptedDocument:
        return ExtractionResult(CVExtractionStatus.FAILED, error=ERROR_ENCRYPTED)
    except Exception:
        # Deliberately broad: pypdf and python-docx raise a wide range of
        # exceptions on malformed input, and none of them should break an upload.
        return ExtractionResult(CVExtractionStatus.FAILED, error=ERROR_PARSE_FAILED)

    text = normalize(raw)
    if len(text) < MIN_TEXT_LENGTH:
        return ExtractionResult(CVExtractionStatus.FAILED, error=ERROR_NO_TEXT_LAYER)
    return ExtractionResult(CVExtractionStatus.COMPLETED, text=text)


def normalize(raw: str) -> str:
    """Collapse the whitespace noise PDF extraction produces and cap the length."""
    text = unicodedata.normalize("NFKC", raw)
    text = _CONTROL_CHARS.sub("", text.replace("\r\n", "\n").replace("\r", "\n"))
    text = _SPACES.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = _BLANK_LINES.sub("\n\n", text).strip()
    return text[:MAX_TEXT_LENGTH]


class _EncryptedDocument(Exception):
    """Password-protected document — a distinct case worth its own user message."""


def _from_pdf(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # Many CVs are "protected" with an empty owner password, which decrypts fine.
        try:
            opened = reader.decrypt("")
        except Exception as exc:
            raise _EncryptedDocument from exc
        if not opened:
            raise _EncryptedDocument
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _from_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    return "\n".join(_docx_blocks(document))


def _docx_blocks(document: DocxDocument) -> list[str]:
    """Paragraphs and table cells in document order.

    Tables matter: a large share of CV templates lay the whole resume out in an
    invisible table, and `document.paragraphs` skips everything inside one.
    """
    blocks: list[str] = []
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            text = Paragraph(child, document).text.strip()
            if text:
                blocks.append(text)
        elif child.tag.endswith("}tbl"):
            for row in Table(child, document).rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Table layout is visual, not semantic — flatten a row to one line.
                line = " | ".join(dict.fromkeys(cell for cell in cells if cell))
                if line:
                    blocks.append(line)
    return blocks
