"""CV text extraction: real PDF and DOCX bytes, no mocking of the parsers."""

import io
import uuid

from docx import Document

from app.cv_versions import extraction

LOREM = (
    "Jane Doe, Senior Backend Engineer. Python, FastAPI, PostgreSQL. "
    "Seven years of experience building payment systems in Berlin and Tbilisi."
)


def _escape(text: str) -> str:
    for char in ("\\", "(", ")"):
        text = text.replace(char, f"\\{char}")
    return text


def _make_pdf(lines: list[str]) -> bytes:
    """Minimal single-page PDF with a real text layer and a correct xref table."""
    shown = "\n".join(f"({_escape(line)}) Tj T*" for line in lines)
    content = f"BT /F1 12 Tf 72 720 Td 14 TL\n{shown}\nET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % number + body + b"\nendobj\n"
    xref_at = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        xref_at,
    )
    return bytes(out)


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload(client, headers, *, data: bytes, filename: str, title: str = "CV"):
    return client.post(
        "/cv-versions/upload",
        headers=headers,
        files={"file": (filename, data, "application/octet-stream")},
        data={"title": title},
    )


def _register_other_user(client):
    email = f"cv-other-{uuid.uuid4().hex[:8]}@test.example"
    response = client.post("/auth/register", json={"email": email, "password": "password123"})
    assert response.status_code == 201, response.text
    response = client.post("/auth/login", json={"email": email, "password": "password123"})
    assert response.status_code == 200, response.text
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


# --- unit level: the extractor itself -------------------------------------


def test_extracts_text_from_pdf():
    result = extraction.extract_text(_make_pdf([LOREM]), ".pdf")

    assert result.status == "completed"
    assert result.error is None
    assert "Senior Backend Engineer" in result.text


def test_extracts_text_from_docx():
    result = extraction.extract_text(_make_docx([LOREM]), ".docx")

    assert result.status == "completed"
    assert "FastAPI" in result.text


def test_reads_docx_laid_out_in_a_table():
    """Many CV templates put the whole resume in a table; paragraphs alone miss it."""
    data = _make_docx(
        ["Jane Doe"],
        table_rows=[
            ["Experience", "Senior Backend Engineer at Acme, seven years of Python"],
            ["Skills", "FastAPI, PostgreSQL, Docker, Kubernetes and message queues"],
        ],
    )

    result = extraction.extract_text(data, ".docx")

    assert result.status == "completed"
    assert "Senior Backend Engineer at Acme" in result.text
    assert "FastAPI, PostgreSQL" in result.text


def test_legacy_doc_is_unsupported_not_failed():
    result = extraction.extract_text(b"\xd0\xcf\x11\xe0 legacy OLE container", ".doc")

    assert result.status == "unsupported"
    assert result.error == extraction.ERROR_UNSUPPORTED_FORMAT


def test_pdf_without_text_layer_reports_no_text_layer():
    """A scan has pages but no extractable characters."""
    result = extraction.extract_text(_make_pdf([]), ".pdf")

    assert result.status == "failed"
    assert result.error == extraction.ERROR_NO_TEXT_LAYER


def test_garbage_bytes_fail_without_raising():
    result = extraction.extract_text(b"this is definitely not a pdf", ".pdf")

    assert result.status == "failed"
    assert result.error == extraction.ERROR_PARSE_FAILED


def test_empty_file_reports_empty():
    result = extraction.extract_text(b"", ".pdf")

    assert result.status == "failed"
    assert result.error == extraction.ERROR_EMPTY_FILE


def test_normalize_collapses_pdf_whitespace_noise():
    normalized = extraction.normalize("  Jane   Doe \r\n\r\n\r\n\r\n  Berlin \x00\x07 ")

    assert normalized == "Jane Doe\n\nBerlin"


def test_normalize_caps_length():
    normalized = extraction.normalize("x" * (extraction.MAX_TEXT_LENGTH + 500))

    assert len(normalized) == extraction.MAX_TEXT_LENGTH


def test_content_hash_tracks_bytes():
    data = _make_pdf([LOREM])

    assert extraction.content_hash(data) == extraction.content_hash(data)
    assert extraction.content_hash(data) != extraction.content_hash(data + b" ")


# --- API level ------------------------------------------------------------


def test_upload_pdf_reports_completed_extraction(client, auth_headers):
    response = _upload(client, auth_headers, data=_make_pdf([LOREM]), filename="cv.pdf")

    assert response.status_code == 201, response.text
    body = response.json()
    assert body["extraction_status"] == "completed"
    assert body["extraction_error"] is None
    assert body["extracted_at"] is not None


def test_upload_legacy_doc_is_stored_but_unsupported(client, auth_headers):
    response = _upload(client, auth_headers, data=b"legacy binary", filename="cv.doc")

    assert response.status_code == 201, response.text
    body = response.json()
    assert body["extraction_status"] == "unsupported"
    assert body["extraction_error"] == "unsupported_format"
    # the file itself is still downloadable — a missing parser must not lose an upload
    assert client.get(f"/cv-versions/{body['id']}/file", headers=auth_headers).status_code == 200


def test_broken_pdf_still_uploads(client, auth_headers):
    """An unparsable document is a status on the row, not a failed request."""
    response = _upload(client, auth_headers, data=b"%PDF-1.4 truncated", filename="cv.pdf")

    assert response.status_code == 201, response.text
    assert response.json()["extraction_status"] == "failed"


def test_extracted_text_is_not_exposed_by_the_api(client, auth_headers):
    created = _upload(client, auth_headers, data=_make_pdf([LOREM]), filename="cv.pdf").json()

    listed = client.get("/cv-versions", headers=auth_headers)
    detail = client.get(f"/cv-versions/{created['id']}", headers=auth_headers)

    assert listed.status_code == 200, listed.text
    assert detail.status_code == 200, detail.text
    assert "extracted_text" not in listed.json()["items"][0]
    assert "extracted_text" not in detail.json()
    assert detail.json()["extraction_status"] == "completed"


def test_retry_endpoint_re_reads_the_stored_file(client, auth_headers):
    created = _upload(client, auth_headers, data=_make_pdf([LOREM]), filename="cv.pdf").json()

    response = client.post(f"/cv-versions/{created['id']}/extract", headers=auth_headers)

    assert response.status_code == 200, response.text
    assert response.json()["extraction_status"] == "completed"


def test_retry_is_scoped_to_the_owner(client, auth_headers):
    created = _upload(client, auth_headers, data=_make_pdf([LOREM]), filename="cv.pdf").json()
    other_headers = _register_other_user(client)

    response = client.post(f"/cv-versions/{created['id']}/extract", headers=other_headers)

    assert response.status_code == 404, response.text
